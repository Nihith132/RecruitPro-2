from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.responses import StreamingResponse
from typing import List
import io
import uuid
import hashlib
from datetime import datetime
import PyPDF2
import docx
from bson import ObjectId

from database.mongodb import get_async_database, get_gridfs
from models.schemas import JobDescription
from llmservices.parser_llm import parse_job_description
from llmservices.topscore_gemini import analyze_multiple_resumes_structured
from routes.auth import verify_firebase_token

router = APIRouter(prefix="/api/jds", tags=["Job Descriptions"])


def generate_file_hash(content: bytes) -> str:
    """Generate SHA256 hash of file content for duplicate detection"""
    return hashlib.sha256(content).hexdigest()


async def extract_text_from_file(file: UploadFile) -> str:
    """Extract text from PDF or DOCX file"""
    content = await file.read()
    
    if file.filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file.filename.endswith('.docx'):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX")


@router.post("/upload", response_model=List[JobDescription])
async def upload_job_descriptions(
    files: List[UploadFile] = File(...),
    uid: str = Depends(verify_firebase_token)
):
    """
    Upload multiple job descriptions with AI parsing and automatic candidate matching
    Prevents duplicate uploads based on file content hash and job title
    """
    db = get_async_database()
    fs = get_gridfs()
    job_descriptions = []
    skipped_files = []
    
    for file in files:
        try:
            # Read file content once
            file_content = await file.read()
            
            # Generate file hash for duplicate detection
            file_hash = generate_file_hash(file_content)
            
            # Check for duplicate file hash
            existing_by_hash = await db.job_descriptions.find_one({
                "uid": uid,
                "file_hash": file_hash
            })
            
            if existing_by_hash:
                print(f"Skipping duplicate JD file (by hash): {file.filename}")
                skipped_files.append({
                    "filename": file.filename,
                    "reason": "Duplicate file - same content already uploaded"
                })
                continue
            
            # Extract text from content
            jd_text = ""
            if file.filename.endswith('.pdf'):
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    jd_text += page.extract_text()
            elif file.filename.endswith('.docx'):
                import docx
                doc = docx.Document(io.BytesIO(file_content))
                jd_text = "\n".join([para.text for para in doc.paragraphs])
            else:
                raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX")
            
            # Parse with AI first to check for duplicate job title
            parsed_data = parse_job_description(jd_text)
            
            # Check for duplicate by job_title and company (if both exist)
            if parsed_data.get('job_title') and parsed_data.get('company'):
                existing_by_title = await db.job_descriptions.find_one({
                    "uid": uid,
                    "job_title": parsed_data['job_title'],
                    "company": parsed_data['company']
                })
                
                if existing_by_title:
                    print(f"Skipping duplicate JD (by title/company): {parsed_data['job_title']} at {parsed_data['company']}")
                    skipped_files.append({
                        "filename": file.filename,
                        "reason": f"Job description '{parsed_data['job_title']}' from '{parsed_data['company']}' already exists"
                    })
                    continue
            
            # Store file in GridFS
            file_id = await fs.upload_from_stream(
                file.filename,
                io.BytesIO(file_content),
                metadata={"content_type": file.content_type, "uploaded_by": uid}
            )
            
            # Create JD document with file hash
            jd_data = {
                "jd_id": str(uuid.uuid4()),
                "uid": uid,
                "file_id": str(file_id),
                "file_hash": file_hash,  # Store hash for duplicate detection
                "jd_filename": file.filename,
                "uploaded_at": datetime.utcnow(),
                **parsed_data
            }
            
            result = await db.job_descriptions.insert_one(jd_data)
            jd_data["_id"] = str(result.inserted_id)
            job_descriptions.append(JobDescription(**jd_data))
            
            # Automatically match with existing candidates
            candidates_cursor = db.candidates.find({"uid": uid})
            candidates = []
            async for candidate in candidates_cursor:
                candidates.append(candidate)
            
            if candidates:
                # Prepare resumes for scoring
                resumes = [{
                    "candidate_id": c["candidate_id"],
                    "name": c.get("name", "Unknown"),
                    "email": c.get("email", ""),
                    "skills": c.get("skills", []),
                    "experience": c.get("experience", ""),  # Fixed: use 'experience' not 'experience_years'
                    "education": c.get("education", []),
                    "certifications": c.get("certifications", [])
                } for c in candidates]
                
                # Analyze with AI
                print(f"Sending {len(resumes)} candidates to AI for analysis during JD upload")  # Debug logging
                scores = analyze_multiple_resumes_structured(jd_text, resumes)
                print(f"Received {len(scores) if isinstance(scores, list) else 'invalid'} scores from AI")  # Debug logging
                
                # Check if scores is valid
                if not scores or not isinstance(scores, list):
                    print(f"Warning: AI analysis failed for JD upload: returned {type(scores).__name__}")
                    scores = []  # Continue without matching
                elif len(scores) == 0:
                    print("Warning: AI analysis returned empty results for JD upload")
                    scores = []  # Continue without matching
                elif len(scores) > 0:
                    # Check first item for errors
                    first_item = scores[0]
                    if isinstance(first_item, dict) and "error" in first_item:
                        print(f"Warning: AI analysis error for JD upload: {first_item.get('error')}")
                        scores = []  # Continue without matching
                
                # Store scores
                for score in scores:
                    # Skip invalid entries
                    if not isinstance(score, dict):
                        print(f"WARNING: Skipping non-dict score entry during JD upload: {type(score)}")
                        continue
                        
                    # Skip error entries
                    if "error" in score:
                        print(f"WARNING: Skipping error entry during JD upload: {score.get('error')}")
                        continue
                        
                    # Compute weighted total score: Skills 50%, Experience 30%, Education 15%, Certs 5%
                    total_score = (
                        score.get("skills_score", 0) * 0.50 +
                        score.get("experience_score", 0) * 0.30 +
                        score.get("education_score", 0) * 0.15 +
                        score.get("certifications_score", 0) * 0.05
                    )
                    
                    score_data = {
                        "score_id": str(uuid.uuid4()),
                        "jd_id": jd_data["jd_id"],
                        "candidate_id": score["candidate_id"],
                        "created_at": datetime.utcnow(),
                        "total_score": total_score,  # Add computed total score
                        **score,  # Spread operator to add all AI fields
                        "uid": uid  # MUST BE LAST - override any uid from AI response
                    }
                    await db.top_scores.insert_one(score_data)
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing {file.filename}: {str(e)}")
    
    # If some files were skipped, log the info
    if skipped_files:
        print(f"Upload complete. Uploaded: {len(job_descriptions)}, Skipped: {len(skipped_files)}")
        for skipped in skipped_files:
            print(f"  - {skipped['filename']}: {skipped['reason']}")
    
    return job_descriptions


@router.get("/", response_model=List[JobDescription])
async def get_job_descriptions(
    skip: int = 0,
    limit: int = 100,
    uid: str = Depends(verify_firebase_token)
):
    """Get all job descriptions with pagination"""
    db = get_async_database()
    cursor = db.job_descriptions.find({"uid": uid}).skip(skip).limit(limit)
    jds = []
    
    async for doc in cursor:
        doc["_id"] = str(doc["_id"])
        jds.append(JobDescription(**doc))
    
    return jds


@router.get("/{jd_id}", response_model=JobDescription)
async def get_job_description(
    jd_id: str,
    uid: str = Depends(verify_firebase_token)
):
    """Get a specific job description by ID"""
    db = get_async_database()
    jd = await db.job_descriptions.find_one({"jd_id": jd_id, "uid": uid})
    
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    jd["_id"] = str(jd["_id"])
    return JobDescription(**jd)


@router.delete("/{jd_id}")
async def delete_job_description(
    jd_id: str,
    uid: str = Depends(verify_firebase_token)
):
    """Delete a job description"""
    db = get_async_database()
    fs = get_gridfs()
    
    # Find JD
    jd = await db.job_descriptions.find_one({"jd_id": jd_id, "uid": uid})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    # Delete file from GridFS
    try:
        file_id = ObjectId(jd["file_id"])
        await fs.delete(file_id)
    except Exception as e:
        print(f"Warning: Could not delete file from GridFS: {e}")
    
    # Delete JD document
    await db.job_descriptions.delete_one({"jd_id": jd_id})
    
    # Delete associated scores
    await db.top_scores.delete_many({"jd_id": jd_id})
    
    return {"message": "Job description deleted successfully"}


@router.get("/download/{jd_id}")
async def download_job_description(
    jd_id: str,
    uid: str = Depends(verify_firebase_token)
):
    """
    Download the original JD file
    Returns the raw file (PDF, DOCX, etc.) for viewing
    """
    db = get_async_database()  # Remove await - it's not async
    fs = get_gridfs()
    
    # Get JD
    jd = await db.job_descriptions.find_one({"jd_id": jd_id})
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    # Get file from GridFS
    try:
        file_id = ObjectId(jd["file_id"])
        grid_out = await fs.open_download_stream(file_id)
        
        # Read file content
        contents = await grid_out.read()
        
        # Get filename and content type
        filename = jd.get("jd_filename", f"{jd.get('job_title', 'job_description')}.pdf")
        
        # Determine content type based on file extension
        if filename.lower().endswith('.pdf'):
            content_type = 'application/pdf'
        elif filename.lower().endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.doc'):
            content_type = 'application/msword'
        else:
            content_type = 'application/octet-stream'
        
        # Return file as streaming response
        return StreamingResponse(
            io.BytesIO(contents),
            media_type=content_type,
            headers={
                'Content-Disposition': f'inline; filename="{filename}"',
                'Content-Type': content_type
            }
        )
    except Exception as e:
        print(f"Error downloading file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to download file: {str(e)}")
