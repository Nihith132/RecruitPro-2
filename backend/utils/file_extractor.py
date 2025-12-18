"""
Enhanced file text extraction with OCR support
Supports: PDF, DOCX, TXT, Images (JPG, PNG), Scanned PDFs
"""

import io
from typing import Union
from fastapi import UploadFile, HTTPException
import PyPDF2
import docx
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes


async def extract_text_from_upload(file: UploadFile) -> str:
    """
    Extract text from various file formats including images
    
    Supported formats:
    - PDF (text-based and scanned)
    - DOCX
    - TXT
    - Images (JPG, PNG, JPEG)
    """
    
    file_content = await file.read()
    filename = file.filename.lower()
    
    try:
        # PDF files
        if filename.endswith('.pdf'):
            return await extract_from_pdf(file_content)
        
        # DOCX files
        elif filename.endswith('.docx'):
            return extract_from_docx(file_content)
        
        # Text files
        elif filename.endswith('.txt'):
            return file_content.decode('utf-8', errors='ignore')
        
        # Image files
        elif filename.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
            return extract_from_image(file_content)
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format. Supported: PDF, DOCX, TXT, JPG, PNG"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error extracting text from file: {str(e)}"
        )


async def extract_from_pdf(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF - handles both text-based and scanned PDFs
    """
    text = ""
    
    try:
        # First try standard text extraction
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # If we got substantial text, return it
        if len(text.strip()) > 100:
            return text
        
        # Otherwise, try OCR (scanned PDF)
        print("PDF appears to be scanned, attempting OCR...")
        images = convert_from_bytes(pdf_bytes)
        
        for i, image in enumerate(images):
            print(f"Processing page {i+1} with OCR...")
            page_text = pytesseract.image_to_string(image, lang='eng')
            text += page_text + "\n"
        
        return text
    
    except Exception as e:
        print(f"PDF extraction error: {e}")
        # Last resort: try OCR
        try:
            images = convert_from_bytes(pdf_bytes)
            text = ""
            for image in images:
                text += pytesseract.image_to_string(image, lang='eng') + "\n"
            return text
        except:
            raise Exception(f"Could not extract text from PDF: {str(e)}")


def extract_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text
        
        return text
    except Exception as e:
        raise Exception(f"Could not extract text from DOCX: {str(e)}")


def extract_from_image(image_bytes: bytes) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(io.BytesIO(image_bytes))
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang='eng')
        
        if not text.strip():
            raise Exception("No text could be extracted from image")
        
        return text
    except Exception as e:
        raise Exception(f"Could not extract text from image: {str(e)}")


def is_supported_file(filename: str) -> bool:
    """Check if file format is supported"""
    filename = filename.lower()
    supported_extensions = ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.bmp', '.tiff']
    return any(filename.endswith(ext) for ext in supported_extensions)
